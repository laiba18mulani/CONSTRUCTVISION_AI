"""Create editable Microsoft Word companions for documentation sources.

Run from the repository root:
    .\.venv\Scripts\python.exe documentation_workspace\render_word.py
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from render_pdfs import ROOT, html_as_markdown, safe_text


def add_inline(paragraph, text: str) -> None:
    """Add a small Markdown subset as editable Word runs."""
    token = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*|\*[^*]+\*)")
    position = 0
    for match in token.finditer(text):
        if match.start() > position:
            paragraph.add_run(safe_text(text[position:match.start()]))
        value = safe_text(match.group(0))
        run = paragraph.add_run(value.strip("`*") )
        if value.startswith("`"):
            run.font.name = "Consolas"
            run.font.size = Pt(8)
        elif value.startswith("**"):
            run.bold = True
        else:
            run.italic = True
        position = match.end()
    if position < len(text):
        paragraph.add_run(safe_text(text[position:]))


def add_table(document: Document, rows: list[list[str]]) -> None:
    columns = max(len(row) for row in rows)
    table = document.add_table(rows=0, cols=columns)
    table.style = "Table Grid"
    for row_index, row in enumerate(rows):
        cells = table.add_row().cells
        for column, value in enumerate(row):
            add_inline(cells[column].paragraphs[0], value)
            if row_index == 0:
                for run in cells[column].paragraphs[0].runs:
                    run.bold = True
    document.add_paragraph()


def configure(document: Document, title: str) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    normal = document.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(6)
    for level, color in ((1, "0B1F3A"), (2, "0B4F6C"), (3, "334155")):
        style = document.styles[f"Heading {level}"]
        style.font.name = "Aptos Display"
        style.font.color.rgb = RGBColor.from_string(color)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run(f"ConstructVision AI | {title}")


def render(markdown: str, output: Path, title: str) -> None:
    document = Document()
    configure(document, title)
    lines = markdown.splitlines()
    index = 0
    in_code = False
    code_lines: list[str] = []

    while index < len(lines):
        line = lines[index]
        stripped = line.strip()
        if stripped.startswith("```"):
            if in_code:
                paragraph = document.add_paragraph()
                paragraph.style = document.styles["Normal"]
                run = paragraph.add_run(safe_text("\n".join(code_lines)))
                run.font.name = "Consolas"
                run.font.size = Pt(8)
                code_lines = []
            in_code = not in_code
            index += 1
            continue
        if in_code:
            code_lines.append(line)
            index += 1
            continue
        if not stripped:
            index += 1
            continue
        if stripped == "---":
            document.add_paragraph("_" * 65)
            index += 1
            continue
        if stripped.startswith("|"):
            rows: list[list[str]] = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                cells = [cell.strip() for cell in lines[index].strip().strip("|").split("|")]
                if not all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells):
                    rows.append(cells)
                index += 1
            if rows:
                add_table(document, rows)
            continue
        heading = re.match(r"^(#{1,4})\s+(.*)$", stripped)
        if heading:
            marks, text = heading.groups()
            level = min(len(marks), 3)
            paragraph = document.add_heading(level=level)
            add_inline(paragraph, text)
            index += 1
            continue
        if re.match(r"^-\s+", stripped):
            paragraph = document.add_paragraph(style="List Bullet")
            add_inline(paragraph, stripped[2:])
            index += 1
            continue
        if re.match(r"^\d+\.\s+", stripped):
            paragraph = document.add_paragraph(style="List Number")
            add_inline(paragraph, re.sub(r"^\d+\.\s+", "", stripped))
            index += 1
            continue
        if stripped.startswith("> "):
            paragraph = document.add_paragraph()
            paragraph.style = document.styles["Quote"]
            add_inline(paragraph, stripped[2:])
            index += 1
            continue
        block = [stripped]
        index += 1
        while index < len(lines):
            candidate = lines[index].strip()
            if not candidate or candidate.startswith(("#", "|", "```", "- ", "> ")) or candidate == "---" or re.match(r"^\d+\.\s+", candidate):
                break
            block.append(candidate)
            index += 1
        paragraph = document.add_paragraph()
        add_inline(paragraph, " ".join(block))

    document.save(output)


def main() -> None:
    for source in sorted(ROOT.glob("*.md")) + sorted(ROOT.glob("*.html")):
        markdown = source.read_text(encoding="utf-8") if source.suffix == ".md" else html_as_markdown(source)
        output = source.with_suffix(".docx")
        if source.suffix == ".html":
            output = source.with_name(f"{source.stem}_from_html.docx")
        render(markdown, output, source.name)
        print(f"Created {output.relative_to(ROOT.parent.parent)}")


if __name__ == "__main__":
    main()
