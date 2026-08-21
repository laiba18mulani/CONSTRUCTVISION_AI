from pathlib import Path
import re

import markdown
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import mm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle, Preformatted
from xml.sax.saxutils import escape

ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "source"
OUT = ROOT / "artifacts"
OUT.mkdir(exist_ok=True)

CSS = """
@page { size: A4; margin: 22mm 18mm 20mm; @bottom-right { content: counter(page); color:#64748b; font-size:9pt; } }
* { box-sizing: border-box; }
body { font-family: Arial, Helvetica, sans-serif; color:#172033; line-height:1.45; font-size:10.5pt; }
h1 { color:#0b1f3a; font-size:28pt; letter-spacing:.02em; margin:0 0 8pt; }
h2 { color:#0b4f6c; border-bottom:1px solid #a7d3df; padding-bottom:5pt; margin-top:24pt; page-break-after:avoid; }
h3 { color:#0b4f6c; margin-top:15pt; page-break-after:avoid; }
h4 { color:#334155; page-break-after:avoid; }
hr { border:0; border-top:2px solid #d8a84e; margin:16pt 0; }
table { width:100%; border-collapse:collapse; margin:10pt 0 14pt; font-size:9pt; }
th { background:#0b4f6c; color:#fff; text-align:left; }
th, td { padding:6pt 7pt; border:1px solid #cbd5e1; vertical-align:top; }
tr:nth-child(even) td { background:#f5f8fa; }
pre { background:#0f172a; color:#e2e8f0; padding:10pt; border-radius:5pt; overflow-wrap:anywhere; white-space:pre-wrap; font-family:Consolas, monospace; font-size:8.5pt; }
code { color:#075985; }
.cover { min-height:235mm; display:flex; flex-direction:column; justify-content:center; text-align:center; page-break-after:always; }
.cover .kicker { color:#d08b18; font-weight:bold; letter-spacing:.16em; text-transform:uppercase; font-size:9pt; }
.cover h1 { font-size:31pt; margin-top:18pt; }
.cover .subtitle { color:#0b4f6c; font-size:15pt; margin-bottom:28pt; }
.meta { background:#eef7fa; border-left:5px solid #0b4f6c; padding:14pt; text-align:left; margin:15pt auto; max-width:145mm; }
.certificate { page-break-after:always; min-height:230mm; }
.placeholder { border:1.5px dashed #94a3b8; background:#f8fafc; padding:18pt; color:#64748b; text-align:center; margin:10pt 0; }
"""


def add_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_font(run, name="Aptos", size=10.5, color="172033", bold=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold


def make_docx(md_text: str, out_path: Path):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)
    for name, size, color in [("Heading 1", 20, "0B1F3A"), ("Heading 2", 15, "0B4F6C"), ("Heading 3", 12, "0B4F6C")]:
        s = styles[name]
        s.font.name = "Aptos Display"
        s.font.size = Pt(size)
        s.font.color.rgb = RGBColor.from_string(color)
        s.font.bold = True

    lines = md_text.splitlines()
    i = 0
    first_title = True
    while i < len(lines):
        line = lines[i]
        if not line.strip():
            i += 1
            continue
        if line.startswith("```"):
            buf = []
            i += 1
            while i < len(lines) and not lines[i].startswith("```"):
                buf.append(lines[i]); i += 1
            p = doc.add_paragraph()
            p.style = doc.styles["Normal"]
            p.paragraph_format.left_indent = Inches(0.15)
            p.paragraph_format.space_before = Pt(4)
            r = p.add_run("\n".join(buf)); set_font(r, "Consolas", 8.5, "334155")
            i += 1; continue
        if line.startswith("|"):
            rows = []
            while i < len(lines) and lines[i].startswith("|"):
                row = [c.strip() for c in lines[i].strip("|").split("|")]
                if not all(set(c) <= set("-: ") for c in row): rows.append(row)
                i += 1
            if rows:
                table = doc.add_table(rows=1, cols=len(rows[0]))
                table.style = "Table Grid"
                for j, val in enumerate(rows[0]):
                    cell = table.rows[0].cells[j]; cell.text = val; add_cell_shading(cell, "0B4F6C")
                    for r in cell.paragraphs[0].runs: set_font(r, "Aptos", 9, "FFFFFF", True)
                for row in rows[1:]:
                    cells = table.add_row().cells
                    for j, val in enumerate(row):
                        cells[j].text = val
                        for r in cells[j].paragraphs[0].runs: set_font(r, "Aptos", 8.5)
            continue
        if line.startswith("# "):
            p = doc.add_paragraph(style="Title")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if first_title else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(line[2:]); set_font(r, "Aptos Display", 24 if first_title else 20, "0B1F3A", True)
            if first_title:
                p.paragraph_format.space_after = Pt(18); first_title = False
            i += 1; continue
        if line.startswith("## ") or line.startswith("### ") or line.startswith("#### "):
            level = len(line) - len(line.lstrip("#"))
            p = doc.add_paragraph(style=f"Heading {min(level-1,3)}")
            r = p.add_run(line[level+1:]); set_font(r, "Aptos Display", {1:20,2:15,3:12}.get(level,11), {1:"0B1F3A",2:"0B4F6C",3:"0B4F6C"}.get(level,"334155"), True)
            i += 1; continue
        if line == "---":
            p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(6)
            r = p.add_run("____________________________________________________________"); set_font(r, "Aptos", 8, "D08B18")
            i += 1; continue
        text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", line)
        text = text.replace("**", "")
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(text); set_font(r)
        i += 1

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("CONSTRUCTVISION AI | Chatake Innoworks Private Limited | MindForgeAI"); set_font(r, "Aptos", 8, "64748B")
    doc.save(out_path)


def make_html(md_text: str, out_path: Path):
    body = markdown.markdown(md_text, extensions=["tables", "fenced_code", "sane_lists"])
    body = body.replace("[Insert Screenshot", '<div class="placeholder">[Insert Screenshot').replace(" here]", " here]</div>")
    html = f"<!doctype html><html><head><meta charset='utf-8'><title>ConstructVision AI Project Report</title><style>{CSS}</style></head><body>{body}</body></html>"
    out_path.write_text(html, encoding="utf-8")
    return html


def make_pdf(md_text: str, out_path: Path):
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="ReportTitle", parent=styles["Title"], fontName="Helvetica-Bold", fontSize=22, leading=27, textColor=colors.HexColor("#0B1F3A"), alignment=TA_CENTER, spaceAfter=18))
    styles.add(ParagraphStyle(name="H1x", parent=styles["Heading1"], fontName="Helvetica-Bold", fontSize=17, leading=21, textColor=colors.HexColor("#0B1F3A"), spaceBefore=16, spaceAfter=8))
    styles.add(ParagraphStyle(name="H2x", parent=styles["Heading2"], fontName="Helvetica-Bold", fontSize=13, leading=16, textColor=colors.HexColor("#0B4F6C"), spaceBefore=12, spaceAfter=6))
    styles.add(ParagraphStyle(name="Bodyx", parent=styles["BodyText"], fontName="Helvetica", fontSize=9.4, leading=13, textColor=colors.HexColor("#172033"), spaceAfter=6))
    styles.add(ParagraphStyle(name="Codex", parent=styles["Code"], fontName="Courier", fontSize=7.2, leading=9, backColor=colors.HexColor("#EEF2F7"), borderPadding=6, spaceBefore=4, spaceAfter=8))
    story = []
    if "## Industry-Grade Project Report" in md_text:
        story.extend([
            Spacer(1, 55*mm),
            Paragraph("CONSTRUCTVISION AI", styles["ReportTitle"]),
            Paragraph("Industry-Grade Project Report", styles["H1x"]),
            Paragraph("An evidence-first civil infrastructure inspection and digital-twin prototype", styles["H2x"]),
            Spacer(1, 10*mm),
            Paragraph("Ritika M. Bhumkar - ritikambhumkar@gmail.com<br/>Laiba Z. Mulani - laiba.mulani.ces.34@gmail.com<br/><br/>Diploma Third Year Civil Engineering<br/><br/><b>Industry Guide</b><br/>Mr. Akash S. Chatake, M.TECH - AIML - BITS PILANI<br/><br/><b>College Guide</b><br/>Ms. Swati P. Maniyal, M.TECH - Structural Engineering<br/><br/><b>Chatake Innoworks Private Limited</b><br/>www.chatakeinnoworks.com<br/><br/><b>MindForgeAI Internship Division</b><br/>https://mindforgeai.co.in<br/><br/>Technical delivery draft | 21 August 2026", styles["Bodyx"]),
            PageBreak(),
        ])
        md_text = md_text.split("---", 1)[1]
    in_code = False; code = []; table_rows = []
    for line in md_text.splitlines():
        if line.startswith("```"):
            if in_code:
                story.append(Preformatted("\n".join(code), styles["Codex"])); code=[]
            in_code = not in_code; continue
        if in_code:
            code.append(line); continue
        if not line.strip():
            if table_rows:
                t = Table(table_rows, repeatRows=1, hAlign="LEFT")
                t.setStyle(TableStyle([("BACKGROUND", (0,0), (-1,0), colors.HexColor("#0B4F6C")), ("TEXTCOLOR", (0,0), (-1,0), colors.white), ("GRID", (0,0), (-1,-1), .35, colors.HexColor("#CBD5E1")), ("VALIGN", (0,0), (-1,-1), "TOP"), ("FONTSIZE", (0,0), (-1,-1), 7.5), ("LEFTPADDING", (0,0), (-1,-1), 5), ("RIGHTPADDING", (0,0), (-1,-1), 5)]))
                story += [t, Spacer(1, 7)]; table_rows=[]
            continue
        if line.startswith("|"):
            cells = [escape(c.strip()) for c in line.strip("|").split("|")]
            if not all(set(c.replace("&lt;", "").replace("&gt;", "")) <= set("-: ") for c in cells):
                table_rows.append([Paragraph(c, styles["Bodyx"]) for c in cells])
            continue
        if line.startswith("# "):
            story += [Paragraph(escape(line[2:]), styles["ReportTitle"]), PageBreak()]
        elif line.startswith("## "):
            story.append(Paragraph(escape(line[3:]), styles["H1x"]))
        elif line.startswith("### "):
            story.append(Paragraph(escape(line[4:]), styles["H2x"]))
        elif line == "---":
            story.append(Spacer(1, 8))
        else:
            text = re.sub(r"\*\*([^*]+)\*\*", r"<b>\1</b>", escape(line))
            story.append(Paragraph(text, styles["Bodyx"]))
    if table_rows:
        t=Table(table_rows, repeatRows=1); t.setStyle(TableStyle([("BACKGROUND", (0,0), (-1,0), colors.HexColor("#0B4F6C")), ("TEXTCOLOR", (0,0), (-1,0), colors.white), ("GRID", (0,0), (-1,-1), .35, colors.HexColor("#CBD5E1")), ("FONTSIZE", (0,0), (-1,-1), 7.5), ("VALIGN", (0,0), (-1,-1), "TOP")])); story.append(t)
    doc = SimpleDocTemplate(str(out_path), pagesize=A4, rightMargin=18*mm, leftMargin=18*mm, topMargin=18*mm, bottomMargin=18*mm)
    doc.build(story)


if __name__ == "__main__":
    report = (SRC / "project_report.md").read_text(encoding="utf-8")
    make_docx(report, OUT / "ConstructVision_AI_Project_Report.docx")
    html = make_html(report, SRC / "project_report.html")
    make_pdf(report, OUT / "ConstructVision_AI_Project_Report.pdf")
    paper = (SRC / "research_paper.md").read_text(encoding="utf-8")
    make_docx(paper, OUT / "ConstructVision_AI_Research_Paper.docx")
    paper_html = markdown.markdown(paper, extensions=["tables", "fenced_code", "sane_lists"])
    make_pdf(paper, OUT / "ConstructVision_AI_Research_Paper.pdf")
